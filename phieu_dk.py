from docx import Document
from docx.shared import Pt, Inches

# Tạo file Word
doc = Document()

# Căn giữa tiêu đề quốc hiệu
p = doc.add_paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
p.alignment = 1
run = p.runs[0]
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph("Độc lập - Tự do - Hạnh phúc")
p.alignment = 1
run = p.runs[0]
run.font.size = Pt(12)

doc.add_paragraph("-----------------------------------------------").alignment = 1

# Thông tin đầu trang
doc.add_paragraph("BV:...........................................................")
doc.add_paragraph("Khoa:.........................................................")

row = doc.add_paragraph()
row.add_run("MS: 01/DKKB-01").bold = True
row.add_run("          Số lưu trữ:..................")
doc.add_paragraph("Mã Y tế: ....../....../.......")

# Tiêu đề
p = doc.add_paragraph("PHIẾU ĐĂNG KÝ KHÁM BỆNH")
p.alignment = 1
p.runs[0].bold = True

# Nội dung form
doc.add_paragraph("- Họ tên người bệnh: ...............................   Tuổi: ........   Nam/Nữ: .......")
doc.add_paragraph("- Dân tộc: ...............................   Nghề nghiệp: ...............................")
doc.add_paragraph("- Mã số BHXH/Thẻ BHYT số: .............................................................")
doc.add_paragraph("- Địa chỉ: ...........................................................................")
doc.add_paragraph("......................................................................................")
doc.add_paragraph("- Số điện thoại liên hệ: .............................................................")
doc.add_paragraph("- Ngày đăng ký khám: ..... giờ ..... phút, ngày ..... tháng ..... năm ......")
doc.add_paragraph("- Triệu chứng chính: ................................................................")
doc.add_paragraph("......................................................................................")
doc.add_paragraph("- Nơi tiếp nhận/Phòng khám: ........................................................")
doc.add_paragraph("- Ghi chú: ..........................................................................")

# Khoảng cách cuối
doc.add_paragraph("\nNgày ..... tháng ..... năm ......")

row = doc.add_paragraph()
row.add_run("Người đăng ký").bold = True
row.add_run(" " * 25)  # tạo khoảng cách
row.add_run("Tiếp nhận đơn vị y tế").bold = True

doc.add_paragraph("(Ký tên, ghi rõ họ tên)                            (Ký tên, đóng dấu)")

# Lưu file
doc.save("phieu_dang_ky_kham_benh.docx")
print("Đã tạo file: phieu_dang_ky_kham_benh.docx")
